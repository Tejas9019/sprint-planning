import io
import csv
import logging
from typing import BinaryIO
import pypdf
import docx

logger = logging.getLogger("app.services.document_parser")

class DocumentParser:
    @staticmethod
    def parse_pdf(file_stream: BinaryIO) -> str:
        logger.info("Parsing PDF file...")
        try:
            reader = pypdf.PdfReader(file_stream)
            text_parts = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to parse PDF document: {str(e)}")

    @staticmethod
    def parse_docx(file_stream: BinaryIO) -> str:
        logger.info("Parsing DOCX file...")
        try:
            doc = docx.Document(file_stream)
            text_parts = [para.text for para in doc.paragraphs]
            # Handle tables inside docx as well
            for table in doc.tables:
                for row in table.rows:
                    text_parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to parse DOCX document: {str(e)}")

    @staticmethod
    def parse_csv(file_stream: BinaryIO) -> str:
        logger.info("Parsing CSV file...")
        try:
            # Read binary stream as text
            content = file_stream.read().decode('utf-8', errors='ignore')
            reader = csv.reader(io.StringIO(content))
            lines = []
            for row in reader:
                lines.append(", ".join(row))
            return "\n".join(lines)
        except Exception as e:
            logger.error(f"Error parsing CSV: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to parse CSV document: {str(e)}")

    @staticmethod
    def parse_txt_or_md(file_stream: BinaryIO) -> str:
        logger.info("Parsing TXT/MD file...")
        try:
            return file_stream.read().decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Error parsing TXT/MD: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to parse text document: {str(e)}")

    @classmethod
    def parse_document(cls, file_name: str, file_stream: BinaryIO) -> str:
        ext = file_name.split('.')[-1].lower() if '.' in file_name else ''
        if ext == 'pdf':
            return cls.parse_pdf(file_stream)
        elif ext == 'docx' or ext == 'doc':
            return cls.parse_docx(file_stream)
        elif ext == 'csv':
            return cls.parse_csv(file_stream)
        elif ext in ['txt', 'md', 'markdown']:
            return cls.parse_txt_or_md(file_stream)
        else:
            # Fallback to general text decoding
            logger.warning(f"Unsupported file extension '{ext}'. Attempting raw text extraction.")
            return cls.parse_txt_or_md(file_stream)
